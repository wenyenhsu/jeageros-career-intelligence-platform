import pytest
from django.urls import reverse
from pathlib import Path

from apps.applications.models import Application
from apps.applications.services.ats_keyword_extractor import AtsKeywordError
from apps.applications.services.cover_letter_tailor_service import (
    CoverLetterTailorService,
)
from apps.applications.services.materials_pack_service import MaterialsPackService


class FakeCoverRewriter:
    def __init__(self, text="Tailored cover letter about Kubernetes.", error=None):
        self.text = text
        self.error = error
        self.calls = []

    def rewrite_cover_letter(
        self, job_title, company, cover_letter_text, job_description
    ):
        self.calls.append(
            {
                "job_title": job_title,
                "company": company,
                "cover_letter_text": cover_letter_text,
                "job_description": job_description,
            }
        )
        if self.error:
            raise AtsKeywordError(self.error)
        return self.text


def _seed_templates(settings):
    root = Path(settings.RESUME_TEMPLATE_ROOT)
    root.mkdir(parents=True, exist_ok=True)
    names = [
        "Cover_Letter_AI.pdf",
        "WenYenHsu_Resume_AI.pdf",
        "Cover_Letter_AI.docx",
        "WenYenHsu_Resume_AI.docx",
        "Cover_Letter_infra.docx",
        "WenYenHsu_Resume_infra.pdf",
        "WenYenHsu_Resume_infra.docx",
    ]
    for name in names:
        (root / name).write_text(f"template {name}", encoding="utf-8")
    return root


@pytest.mark.django_db
def test_tailor_rewrites_copied_cover_letter_and_backs_up_original(
    application, job, application_materials_root, settings
):
    _seed_templates(settings)
    job.source_url = "https://example.com/jobs/backend"
    job.save(update_fields=["source_url"])
    MaterialsPackService().apply_pack(application, "AI")
    rewriter = FakeCoverRewriter()
    source_resume = Path(settings.RESUME_TEMPLATE_ROOT) / "WenYenHsu_Resume_AI.pdf"

    result = CoverLetterTailorService(
        extractor=rewriter,
        fetch_description=lambda url: "Need Kubernetes and Python.",
    ).run(application)

    dest = Path(application.materials_local_path)
    backup = dest / "Cover_Letter_AI.original.pdf"
    assert result["tailored"] is True
    assert backup.read_text(encoding="utf-8") == "template Cover_Letter_AI.pdf"
    assert "Kubernetes" in CoverLetterTailorService._read_text(
        dest / "Cover_Letter_AI.pdf"
    )
    assert (dest / "WenYenHsu_Resume_AI.pdf").read_text(encoding="utf-8") == (
        "template WenYenHsu_Resume_AI.pdf"
    )
    assert source_resume.read_text(encoding="utf-8") == "template WenYenHsu_Resume_AI.pdf"
    assert rewriter.calls[0]["company"] == job.company.name
    assert "Kubernetes" in rewriter.calls[0]["job_description"]
    job.refresh_from_db()
    assert job.description == "Need Kubernetes and Python."


@pytest.mark.django_db
def test_tailor_skips_without_url_or_description(
    application, application_materials_root, settings
):
    _seed_templates(settings)
    MaterialsPackService().apply_pack(application, "AI")
    dest = Path(application.materials_local_path)
    original = (dest / "Cover_Letter_AI.pdf").read_text(encoding="utf-8")

    result = CoverLetterTailorService(
        extractor=FakeCoverRewriter(),
        fetch_description=lambda url: "",
    ).run(application)

    assert result["skipped"] is True
    assert (dest / "Cover_Letter_AI.pdf").read_text(encoding="utf-8") == original
    assert not (dest / "Cover_Letter_AI.original.pdf").exists()


@pytest.mark.django_db
def test_tailor_uses_stored_description_when_url_fetch_fails(
    application, job, application_materials_root, settings
):
    _seed_templates(settings)
    job.source_url = "https://example.com/jobs/backend"
    job.description = "Need Terraform."
    job.save(update_fields=["source_url", "description"])
    MaterialsPackService().apply_pack(application, "AI")
    rewriter = FakeCoverRewriter(text="Tailored cover letter about Terraform.")

    result = CoverLetterTailorService(
        extractor=rewriter,
        fetch_description=lambda url: "",
    ).run(application)

    assert result["tailored"] is True
    assert rewriter.calls[0]["job_description"] == "Need Terraform."


@pytest.mark.django_db
def test_tailor_leaves_copy_when_llm_fails(
    application, job, application_materials_root, settings
):
    _seed_templates(settings)
    job.source_url = "https://example.com/jobs/backend"
    job.save(update_fields=["source_url"])
    MaterialsPackService().apply_pack(application, "AI")
    dest = Path(application.materials_local_path)
    original = (dest / "Cover_Letter_AI.pdf").read_text(encoding="utf-8")

    result = CoverLetterTailorService(
        extractor=FakeCoverRewriter(error="Ollama down"),
        fetch_description=lambda url: "Need Python.",
    ).run(application)

    assert result["tailored"] is False
    assert "unchanged" in result["error"]
    assert (dest / "Cover_Letter_AI.pdf").read_text(encoding="utf-8") == original
    assert not (dest / "Cover_Letter_AI.original.pdf").exists()


@pytest.mark.django_db
def test_create_view_tailors_cover_letter_from_job_url(
    client, user, application_materials_root, settings, monkeypatch
):
    _seed_templates(settings)
    rewriter = FakeCoverRewriter()
    monkeypatch.setattr(
        "apps.applications.services.cover_letter_tailor_service.CoverLetterTailorService",
        lambda: CoverLetterTailorService(
            extractor=rewriter,
            fetch_description=lambda url: "Need Kubernetes.",
        ),
    )

    response = client.post(
        reverse("application-create"),
        data={
            "user": user.pk,
            "status": Application.Status.SAVED,
            "priority": 3,
            "company": "Rippling",
            "job_title": "Software Engineer",
            "source_url": "https://example.com/jobs/swe",
            "materials_pack": "AI",
        },
        follow=True,
    )

    assert response.status_code == 200
    application = Application.objects.get(user=user)
    dest = Path(application.materials_local_path)
    assert "Kubernetes" in CoverLetterTailorService._read_text(
        dest / "Cover_Letter_AI.pdf"
    )
    assert (dest / "Cover_Letter_AI.original.pdf").read_text(encoding="utf-8") == (
        "template Cover_Letter_AI.pdf"
    )
    application.job_post.refresh_from_db()
    assert application.job_post.description == "Need Kubernetes."


@pytest.mark.django_db
def test_cover_letter_tailor_does_not_import_market_fit():
    from apps.applications.services import cover_letter_tailor_service

    combined = Path(cover_letter_tailor_service.__file__).read_text(encoding="utf-8")
    assert "MarketFit" not in combined
    assert "ResumeGap" not in combined
    assert "SkillDemand" not in combined
    assert "SkillPipeline" not in combined


@pytest.mark.django_db
def test_create_page_includes_cover_letter_progress_widget(client):
    response = client.get(reverse("application-create"))

    assert response.status_code == 200
    content = response.content.decode()
    assert 'id="globalCoverLetterStatus"' in content
    assert "data-cover-letter-progress-form" in content
    assert "jaegerosStartCoverLetterRun" in content


@pytest.mark.django_db
def test_cover_letter_status_endpoint_returns_progress(client):
    from apps.applications.services.cover_letter_tailor_service import (
        save_cover_letter_progress,
    )

    save_cover_letter_progress(
        "abc123",
        status="STARTED",
        progress=40,
        label="Fetching job description",
    )
    response = client.get(
        reverse("application-cover-letter-status"),
        {"run_id": "abc123"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "STARTED"
    assert payload["progress"] == 40
    assert payload["current_step"]["label"] == "Fetching job description"


@pytest.mark.django_db
def test_create_json_returns_cover_letter_run_id(
    client, user, application_materials_root, settings, monkeypatch
):
    _seed_templates(settings)
    monkeypatch.setattr(
        "apps.applications.services.cover_letter_tailor_service.CoverLetterTailorService",
        lambda: CoverLetterTailorService(
            extractor=FakeCoverRewriter(),
            fetch_description=lambda url: "Need Kubernetes.",
        ),
    )

    response = client.post(
        reverse("application-create"),
        data={
            "user": user.pk,
            "status": Application.Status.SAVED,
            "priority": 3,
            "company": "Rippling",
            "job_title": "Software Engineer",
            "source_url": "https://example.com/jobs/swe",
            "materials_pack": "AI",
            "cover_letter_run_id": "run123xyz",
        },
        HTTP_ACCEPT="application/json",
        HTTP_X_REQUESTED_WITH="XMLHttpRequest",
    )

    assert response.status_code == 202
    payload = response.json()
    assert payload["success"] is True
    assert payload["cover_letter_run_id"] == "run123xyz"
    status = client.get(
        reverse("application-cover-letter-status"),
        {"run_id": "run123xyz"},
    ).json()
    assert status["status"] == "SUCCESS"
    assert status["progress"] == 100


def _write_letter_docx(path, company="Rippling"):
    from docx import Document
    from docx.shared import Pt

    document = Document()
    name = document.add_paragraph("WEN-YEN (HANK) HSU")
    name.runs[0].bold = True
    name.runs[0].font.size = Pt(14)
    document.add_paragraph(
        "Los Angeles, CA | +1 (213) 536-3478 | godhanko@gmail.com"
    )
    document.add_paragraph("August 19, 2026")
    document.add_paragraph("")
    document.add_paragraph(company)
    document.add_paragraph("Machine Learning Team")
    document.add_paragraph("San Francisco, CA")
    document.add_paragraph(f"Dear {company} Hiring Team:")
    document.add_paragraph(
        f"I am writing to apply for the intern position at {company}."
    )
    document.add_paragraph("My research focuses on multimodal AI agents.")
    document.add_paragraph("Previous experience at TSMC and Entegris.")
    document.add_paragraph("I would welcome the opportunity to contribute.")
    document.add_paragraph("Best regards,")
    document.add_paragraph("Wen-Yen (Hank) Hsu")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


@pytest.mark.django_db
def test_tailor_keeps_cover_letter_layout_after_jd_rewrite(
    application, job, application_materials_root, settings
):
    from docx import Document
    from pypdf import PdfReader

    job.company.name = "Crowe"
    job.company.save(update_fields=["name"])
    job.source_url = "https://example.com/jobs/backend"
    job.save(update_fields=["source_url"])
    root = Path(settings.RESUME_TEMPLATE_ROOT)
    _write_letter_docx(root / "Cover_Letter_AI.docx")
    (root / "Cover_Letter_AI.pdf").write_text("template pdf", encoding="utf-8")
    (root / "WenYenHsu_Resume_AI.pdf").write_text("resume", encoding="utf-8")
    MaterialsPackService().apply_pack(application, "AI")
    mashed = (
        "Cover Letter HANK HSU Los Angeles, CA | other@example.com "
        "August 19, 2026 Crowe AI Team San Francisco, CA "
        "Dear Crowe Hiring Team: I am writing about Kubernetes. "
        "Best regards, Someone Else"
    )
    rewriter = FakeCoverRewriter(text=mashed)

    result = CoverLetterTailorService(
        extractor=rewriter,
        fetch_description=lambda url: "<p>Need <b>Kubernetes</b>.</p>",
    ).run(application)

    dest = Path(application.materials_local_path)
    doc = Document(dest / "Cover_Letter_AI.docx")
    texts = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    ]
    pdf_text = PdfReader(str(dest / "Cover_Letter_AI.pdf")).pages[0].extract_text()

    assert result["tailored"] is True
    assert "\n" in rewriter.calls[0]["cover_letter_text"]
    assert "<p>" not in rewriter.calls[0]["job_description"]
    assert texts[0] == "WEN-YEN (HANK) HSU"
    assert texts[1].startswith("Los Angeles, CA")
    assert "godhanko@gmail.com" in texts[1]
    assert any(line.startswith("Dear Crowe") for line in texts)
    assert any(line == "San Francisco, CA" for line in texts)
    assert any("Kubernetes" in line for line in texts)
    assert "Cover Letter" not in texts
    assert doc.paragraphs[0].runs[0].bold is True
    assert pdf_text.splitlines()[0].strip() == "WEN-YEN (HANK) HSU"
    assert "Cover Letter" not in pdf_text.splitlines()[0]
    assert "Kubernetes" in pdf_text
    assert "Rippling" not in pdf_text


@pytest.mark.django_db
def test_tailor_updates_pdf_when_model_copies_original_letter(
    application, job, application_materials_root, settings
):
    from docx import Document
    from pypdf import PdfReader

    job.company.name = "Crowe"
    job.company.save(update_fields=["name"])
    job.title = "AI Functional Intern"
    job.source_url = "https://example.com/jobs/backend"
    job.save(update_fields=["title", "source_url"])
    root = Path(settings.RESUME_TEMPLATE_ROOT)
    _write_letter_docx(root / "Cover_Letter_AI.docx")
    (root / "Cover_Letter_AI.pdf").write_text("template pdf", encoding="utf-8")
    (root / "WenYenHsu_Resume_AI.pdf").write_text("resume", encoding="utf-8")
    MaterialsPackService().apply_pack(application, "AI")

    class EchoRewriter:
        def rewrite_cover_letter(
            self, job_title, company, cover_letter_text, job_description
        ):
            return cover_letter_text

    result = CoverLetterTailorService(
        extractor=EchoRewriter(),
        fetch_description=lambda url: "Need Kubernetes.",
    ).run(application)

    dest = Path(application.materials_local_path)
    pdf_text = PdfReader(str(dest / "Cover_Letter_AI.pdf")).pages[0].extract_text()
    doc = Document(dest / "Cover_Letter_AI.docx")
    texts = [paragraph.text for paragraph in doc.paragraphs]

    assert result["tailored"] is True
    assert "Crowe" in pdf_text
    assert "Rippling" not in pdf_text
    assert "AI Functional Intern" in pdf_text
    assert any(line == "Crowe" for line in texts)
    assert any("Dear Crowe" in line for line in texts)
    assert "Rippling" not in "\n".join(texts)
