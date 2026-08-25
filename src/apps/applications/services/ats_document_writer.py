from copy import deepcopy
from pathlib import Path
import shutil
import textwrap

from .cover_letter_layout import letter_slot_indexes, normalize_letter_text


class AtsDocumentWriter:
    """Write Simplify/ATS-like single-column Resume.pdf and Cover-Letter.pdf."""

    resume_name = "Resume.pdf"
    cover_letter_name = "Cover-Letter.pdf"
    original_suffix = ".original.pdf"

    def write_resume(self, folder, text, source_path=None):
        return self._write_named_pdf(
            folder,
            filename=self.resume_name,
            title="Resume",
            text=text,
            source_path=source_path,
        )

    def write_cover_letter(self, folder, text, source_path=None):
        return self._write_named_pdf(
            folder,
            filename=self.cover_letter_name,
            title="Cover Letter",
            text=text,
            source_path=source_path,
        )

    def _write_named_pdf(self, folder, filename, title, text, source_path=None):
        folder = Path(folder)
        folder.mkdir(parents=True, exist_ok=True)
        target = folder / filename
        source = Path(source_path) if source_path else None
        backup = ""
        if (
            source is not None
            and source.exists()
            and source.resolve() == target.resolve()
        ):
            backup_path = folder / f"{target.stem}{self.original_suffix}"
            if not backup_path.exists():
                shutil.copy2(source, backup_path)
            backup = backup_path.name
        self.write_plain_pdf(target, title=title, body=text)
        self._write_plain_docx(target.with_suffix(".docx"), title=title, body=text)
        return {"path": target, "backup": backup}

    def overwrite_file(self, path, title, text):
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        backup_path = path.with_name(f"{path.stem}.original{path.suffix}")
        backup = ""
        if path.exists() and not backup_path.exists():
            shutil.copy2(path, backup_path)
            backup = backup_path.name
        suffix = path.suffix.casefold()
        if suffix == ".pdf":
            self.write_plain_pdf(path, title=title, body=text)
        elif suffix == ".docx":
            self._write_plain_docx(path, title=title, body=text)
        else:
            path.write_text(str(text or "").rstrip() + "\n", encoding="utf-8")
        return {"path": path, "backup": backup}

    def overwrite_cover_letter(self, path, text):
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        backup_path = path.with_name(f"{path.stem}.original{path.suffix}")
        backup = ""
        if path.exists() and not backup_path.exists():
            shutil.copy2(path, backup_path)
            backup = backup_path.name
        body = normalize_letter_text(text)
        suffix = path.suffix.casefold()
        if suffix == ".docx":
            source = backup_path if backup_path.exists() else path
            if not self._rewrite_docx_preserving_layout(source, path, body):
                self._write_plain_docx(path, title="", body=body)
        elif suffix == ".pdf":
            self.write_plain_pdf(path, title="", body=body)
        else:
            path.write_text(body.rstrip() + "\n", encoding="utf-8")
        return {"path": path, "backup": backup}

    @classmethod
    def write_plain_pdf(cls, path, title, body):
        chunks = []
        heading = str(title or "").strip()
        if heading:
            chunks.append(heading)
            chunks.append("")
        chunks.append(str(body or "").replace("\r\n", "\n").strip())
        lines = cls._wrap_lines("\n".join(chunks), width=72)
        commands = ["BT /F1 11 Tf 72 720 Td 16 TL"]
        commands.extend(f"({cls._pdf_escape(line)}) Tj T*" for line in lines)
        commands.append("ET")
        content_bytes = ("\n".join(commands) + "\n").encode("latin-1")
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            (
                b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
            ),
            b"<< /Length %d >>\nstream\n" % len(content_bytes)
            + content_bytes
            + b"endstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]
        output = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, obj in enumerate(objects, start=1):
            offsets.append(len(output))
            output.extend(f"{index} 0 obj\n".encode("ascii"))
            output.extend(obj)
            output.extend(b"\nendobj\n")
        xref = len(output)
        output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        output.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        output.extend(
            (
                f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref}\n%%EOF\n"
            ).encode("ascii")
        )
        Path(path).write_bytes(bytes(output))

    @classmethod
    def _rewrite_docx_preserving_layout(cls, source, target, text):
        from docx import Document

        from .cover_letter_layout import parse_letter

        source = Path(source)
        if source.suffix.casefold() != ".docx" or not source.is_file():
            return False
        try:
            document = Document(str(source))
        except Exception:
            return False
        paragraphs = list(document.paragraphs)
        slots = letter_slot_indexes(paragraph.text for paragraph in paragraphs)
        if slots["greeting"] is None or not slots["body"]:
            return False

        parts = parse_letter(text)
        if parts["greeting"]:
            cls._set_paragraph_text(paragraphs[slots["greeting"]], parts["greeting"])
        if parts["closing"] and slots["closing"] is not None:
            cls._set_paragraph_text(paragraphs[slots["closing"]], parts["closing"])
        cls._replace_slot_paragraphs(paragraphs, slots["body"], parts["body"])
        cls._replace_slot_paragraphs(
            paragraphs, slots["recipient"], parts["recipient"]
        )
        Path(target).parent.mkdir(parents=True, exist_ok=True)
        document.save(target)
        return True

    @classmethod
    def _replace_slot_paragraphs(cls, paragraphs, indexes, values):
        indexes = list(indexes or [])
        values = [str(value).strip() for value in (values or []) if str(value).strip()]
        if not indexes or not values:
            return
        for index, value in zip(indexes, values):
            cls._set_paragraph_text(paragraphs[index], value)
        if len(values) < len(indexes):
            for index in reversed(indexes[len(values) :]):
                cls._delete_paragraph(paragraphs[index])
            return
        if len(values) <= len(indexes):
            return
        last = paragraphs[indexes[-1]]
        for value in values[len(indexes) :]:
            last = cls._clone_paragraph_after(last, value)

    @staticmethod
    def _clone_paragraph_after(paragraph, text):
        from docx.text.paragraph import Paragraph

        new_p = deepcopy(paragraph._p)
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        AtsDocumentWriter._set_paragraph_text(new_para, text)
        return new_para

    @staticmethod
    def _delete_paragraph(paragraph):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    @staticmethod
    def _set_paragraph_text(paragraph, text):
        value = str(text or "")
        runs = paragraph.runs
        if not runs:
            paragraph.add_run(value)
            return
        runs[0].text = value
        for run in runs[1:]:
            run.text = ""

    @staticmethod
    def _write_plain_docx(path, title, body):
        from docx import Document
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Inches, Pt

        document = Document()
        for section in document.sections:
            section.top_margin = Inches(0.65)
            section.bottom_margin = Inches(0.65)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        heading = str(title or "").strip()
        if heading:
            heading_para = document.add_paragraph(heading)
            heading_para.runs[0].bold = True
            heading_para.runs[0].font.size = Pt(14)
            heading_para.runs[0].font.name = "Calibri"
        lines = str(body or "").replace("\r\n", "\n").split("\n") or [""]
        for index, block in enumerate(lines):
            paragraph = document.add_paragraph(block)
            is_body = len(block) > 80
            paragraph.paragraph_format.space_after = Pt(8 if is_body else 0)
            paragraph.paragraph_format.space_before = Pt(8 if is_body else 0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if not paragraph.runs:
                continue
            run = paragraph.runs[0]
            run.font.name = "Calibri"
            if index == 0 and not heading and block.isupper() and len(block) < 80:
                run.bold = True
                run.font.size = Pt(14)
            elif "|" in block or "@" in block:
                run.font.size = Pt(9.5)
            else:
                run.font.size = Pt(11)
                if is_body:
                    paragraph.paragraph_format.line_spacing = 1.15
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        document.save(path)

    @classmethod
    def _wrap_lines(cls, text, width=72):
        lines = []
        for raw in str(text or "").replace("\r\n", "\n").split("\n"):
            chunk = cls._latin1(raw)
            if not chunk.strip():
                lines.append("")
                continue
            lines.extend(textwrap.wrap(chunk, width=width) or [""])
        return lines[:48] or [""]

    @staticmethod
    def _latin1(value):
        replacements = {
            "\u2018": "'",
            "\u2019": "'",
            "\u201c": '"',
            "\u201d": '"',
            "\u2013": "-",
            "\u2014": "-",
            "\u00a0": " ",
        }
        text = str(value or "")
        for source, target in replacements.items():
            text = text.replace(source, target)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    @staticmethod
    def _pdf_escape(value):
        return (
            str(value or "")
            .replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )
